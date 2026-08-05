from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date
from io import BytesIO

import streamlit as st
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.page import PageMargins

try:
    from .ui_components import page_header
except ImportError:  # 단독 실행·테스트용
    def page_header(_section: str, title: str, description: str, _code: str) -> None:
        st.title(title)
        st.caption(description)


APP_TITLE = "보험 리모델링 비교 제안서"

NAVY = "17365D"
NAVY2 = "245889"
GOLD = "C9A24D"
GOLD_LIGHT = "FFF4D6"
BLUE_LIGHT = "EAF2F8"
GREEN_LIGHT = "E8F4F2"
SOFT = "F7F9FC"
WHITE = "FFFFFF"
INK = "25364A"
MUTED = "667085"
GREEN = "24745A"
RED = "C00000"
LINE = "B8C2CF"
THIN = Side(style="thin", color=LINE)

PAYMENT_TYPES = ["확정 납입", "갱신형", "종신납", "계속 납입"]
CONTRACT_ACTIONS = ["유지", "감액", "일부 특약 조정", "해지", "신규 승인 후 결정", "추가 확인"]


@dataclass
class NewPlan:
    name: str = ""
    monthly: int = 0
    payment_type: str = "확정 납입"
    years: int = 20
    custom_months: int = 0

    @property
    def months(self) -> int:
        if self.payment_type != "확정 납입":
            return 0
        return self.custom_months if self.custom_months > 0 else self.years * 12

    @property
    def fixed_total(self) -> int:
        return self.monthly * self.months


@dataclass
class ExistingContract:
    company: str = ""
    product: str = ""
    action: str = "유지"
    detail: str = ""


@dataclass
class Person:
    name: str
    old_monthly: int
    old_total: int
    retained_monthly: int
    retained_total: int
    plans: list[NewPlan] = field(default_factory=list)
    coverage: str = ""
    contracts: list[ExistingContract] = field(default_factory=list)

    @property
    def new_plan_monthly(self) -> int:
        return sum(p.monthly for p in self.plans)

    @property
    def after_monthly(self) -> int:
        return self.retained_monthly + self.new_plan_monthly

    @property
    def after_total(self) -> int:
        return self.retained_total + sum(p.fixed_total for p in self.plans)

    @property
    def monthly_change(self) -> int:
        return self.after_monthly - self.old_monthly

    @property
    def total_change(self) -> int:
        return self.after_total - self.old_total

    @property
    def excluded_plans(self) -> list[NewPlan]:
        return [p for p in self.plans if p.payment_type != "확정 납입"]


def money(value: object) -> int:
    text = re.sub(r"[^0-9-]", "", str(value or ""))
    try:
        return int(text) if text not in {"", "-"} else 0
    except ValueError:
        return 0


def clean(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def safe_filename(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", value).strip() or "보험리모델링_비교안"


def won(value: int) -> str:
    return f"{int(value):,}원"


def rate(old: int, new: int) -> float | None:
    return (old - new) / old if old > new and old else None


def change_amount(old: int, new: int) -> str:
    if old > new:
        return f"{old-new:,}원 절감"
    if old < new:
        return f"{new-old:,}원 증가"
    return "변동 없음"


def change_rate(old: int, new: int) -> str:
    r = rate(old, new)
    return f"{r:.1%} 감소" if r is not None else ""


def combined(people: list[Person]) -> dict[str, int]:
    return {
        "old_monthly": sum(p.old_monthly for p in people),
        "after_monthly": sum(p.after_monthly for p in people),
        "old_total": sum(p.old_total for p in people),
        "after_total": sum(p.after_total for p in people),
    }


def _state_count(key: str, default: int) -> int:
    st.session_state.setdefault(key, default)
    return int(st.session_state[key])


def _money_input(label: str, key: str, help_text: str | None = None) -> int:
    raw = st.text_input(label, key=key, placeholder="예: 694,580", help=help_text)
    value = money(raw)
    if raw:
        st.caption(f"{value:,}원")
    return value


def render_plan_inputs(person_no: int) -> list[NewPlan]:
    count_key = f"rm_plan_count_{person_no}"
    count = _state_count(count_key, 4)
    plans: list[NewPlan] = []
    for i in range(count):
        title = clean(st.session_state.get(f"rm_plan_name_{person_no}_{i}")) or f"신규 보험 {i+1}"
        with st.expander(title, expanded=i < 2):
            c1, c2 = st.columns([2.2, 1])
            with c1:
                name = st.text_input("보험 또는 보장 구성명", key=f"rm_plan_name_{person_no}_{i}", placeholder="예: 암·뇌·심장 진단비")
            with c2:
                premium = st.text_input("월 보험료", key=f"rm_plan_premium_{person_no}_{i}", placeholder="예: 128,589")
            c3, c4, c5 = st.columns(3)
            with c3:
                ptype = st.selectbox("납입 유형", PAYMENT_TYPES, key=f"rm_plan_type_{person_no}_{i}")
            years = 20
            custom = False
            months = 0
            if ptype == "확정 납입":
                with c4:
                    years = st.selectbox("납입기간", [5, 10, 15, 20, 25, 30], index=3, format_func=lambda x: f"{x}년", key=f"rm_plan_years_{person_no}_{i}")
                with c5:
                    custom = st.checkbox("개월 수 직접 입력", key=f"rm_plan_custom_on_{person_no}_{i}")
                if custom:
                    months = int(st.number_input("직접 입력할 납입 개월 수", min_value=1, max_value=1200, value=240, step=1, key=f"rm_plan_months_{person_no}_{i}"))
            plans.append(NewPlan(clean(name), money(premium), ptype, int(years), int(months)))
    a, b = st.columns(2)
    with a:
        if st.button("＋ 신규 보험 추가", key=f"rm_plan_add_{person_no}", use_container_width=True):
            st.session_state[count_key] = min(12, count + 1)
            st.rerun()
    with b:
        if st.button("－ 마지막 신규 보험 삭제", key=f"rm_plan_remove_{person_no}", disabled=count <= 1, use_container_width=True):
            st.session_state[count_key] = max(1, count - 1)
            st.rerun()
    return [p for p in plans if p.name or p.monthly]


def render_contract_inputs(person_no: int) -> list[ExistingContract]:
    count_key = f"rm_contract_count_{person_no}"
    count = _state_count(count_key, 2)
    result: list[ExistingContract] = []
    for i in range(count):
        with st.expander(f"기존 계약 {i+1}", expanded=i == 0):
            a, b = st.columns(2)
            with a:
                company = st.text_input("보험회사", key=f"rm_contract_company_{person_no}_{i}")
                product = st.text_input("상품명", key=f"rm_contract_product_{person_no}_{i}")
            with b:
                action = st.selectbox("처리 방향", CONTRACT_ACTIONS, key=f"rm_contract_action_{person_no}_{i}")
                detail = st.text_input("변경 내용", key=f"rm_contract_detail_{person_no}_{i}")
            result.append(ExistingContract(clean(company), clean(product), action, clean(detail)))
    a, b = st.columns(2)
    with a:
        if st.button("＋ 기존 계약 추가", key=f"rm_contract_add_{person_no}", use_container_width=True):
            st.session_state[count_key] = min(12, count + 1)
            st.rerun()
    with b:
        if st.button("－ 마지막 기존 계약 삭제", key=f"rm_contract_remove_{person_no}", disabled=count <= 1, use_container_width=True):
            st.session_state[count_key] = max(1, count - 1)
            st.rerun()
    return [c for c in result if c.company or c.product or c.detail]


def render_person_inputs(person_no: int, detailed: bool) -> Person:
    st.subheader(f"고객 {person_no}")
    name = st.text_input("고객명", key=f"rm_name_{person_no}", placeholder="예: 홍길동")
    a, b = st.columns(2)
    with a:
        old_monthly = _money_input("기존 월 보험료", f"rm_old_monthly_{person_no}")
        old_total = _money_input("기존 납입 예정 총액", f"rm_old_total_{person_no}")
    with b:
        retained_monthly = _money_input("유지하는 기존 보험료", f"rm_retained_monthly_{person_no}")
        retained_total = _money_input("유지 보험의 남은 확정 납입 예정 총액", f"rm_retained_total_{person_no}", "사용자가 확인한 합계 금액을 직접 입력합니다.")
    st.markdown("#### 새롭게 가입하는 보험")
    plans = render_plan_inputs(person_no)
    coverage = st.text_area("새롭게 확보되는 핵심 보장", key=f"rm_coverage_{person_no}", placeholder="예: 암·뇌·심장 진단비 보완 · 주요 치료비 강화", height=75)
    contracts = []
    if detailed:
        st.markdown("#### 기존 계약별 유지·감액·해지")
        contracts = render_contract_inputs(person_no)
    person = Person(clean(name), old_monthly, old_total, retained_monthly, retained_total, plans, clean(coverage), contracts)
    st.info(
        f"신규 보험료 합계 {person.new_plan_monthly:,}원  ·  "
        f"변경 후 월 보험료 {person.after_monthly:,}원  ·  "
        f"변경 후 납입 예정 총액 {person.after_total:,}원"
    )
    if person.excluded_plans:
        names = ", ".join(p.name or "미입력 보험" for p in person.excluded_plans)
        st.caption(f"총납입액 비교 제외: {names}")
    return person


# ---------------- Excel ----------------
def _border() -> Border:
    return Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def _merge(ws, address: str, value: object, *, fill: str | None = None, color: str = INK,
           size: float = 10, bold: bool = False, border: bool = True) -> None:
    ws.merge_cells(address)
    cell = ws[address.split(":")[0]]
    cell.value = value
    cell.font = Font(name="맑은 고딕", size=size, bold=bold, color=color)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True, shrink_to_fit=True)
    if fill:
        cell.fill = PatternFill("solid", fgColor=fill)
    if border:
        for row in ws[address]:
            for item in row:
                item.border = _border()


def _excel_setup(ws, last_col: str, last_row: int) -> None:
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins = PageMargins(left=.25, right=.25, top=.30, bottom=.30, header=.1, footer=.1)
    ws.print_options.horizontalCentered = True
    ws.print_area = f"A1:{last_col}{last_row}"


def _excel_top(ws, people: list[Person], title: str) -> None:
    totals = combined(people)
    _merge(ws, "A1:R2", title, color=NAVY, size=19, bold=True, border=False)
    subtitle = "보험료 부담과 필요한 보장을 함께 비교해 오래 유지하기 쉬운 구조로 재구성했습니다."
    _merge(ws, "A3:R3", subtitle, color=MUTED, size=10, bold=True, border=False)
    labels = (["월 보험료 변화", "변경 후 월 보험료", "납입 예정 총액 변화"] if len(people) == 1 else
              ["합산 월 보험료 변화", "변경 후 합산 월 보험료", "납입 예정 총액 변화"])
    for address, label in zip(("A5:F5", "G5:L5", "M5:R5"), labels):
        _merge(ws, address, label, fill=NAVY2, color=WHITE, size=9, bold=True)
    _merge(ws, "A6:D6", change_amount(totals["old_monthly"], totals["after_monthly"]), fill=GOLD_LIGHT, color=NAVY, size=11, bold=True)
    _merge(ws, "E6:F6", change_rate(totals["old_monthly"], totals["after_monthly"]), fill=GOLD_LIGHT, color=RED, size=11, bold=True)
    _merge(ws, "G6:L6", won(totals["after_monthly"]), fill=GOLD_LIGHT, color=NAVY, size=11, bold=True)
    _merge(ws, "M6:P6", change_amount(totals["old_total"], totals["after_total"]), fill=GOLD_LIGHT, color=NAVY, size=11, bold=True)
    _merge(ws, "Q6:R6", change_rate(totals["old_total"], totals["after_total"]), fill=GOLD_LIGHT, color=RED, size=11, bold=True)


def _excel_person_panel(ws, p: Person, left: int, right: int, top: int, max_plans: int) -> int:
    L, R = get_column_letter(left), get_column_letter(right)
    _merge(ws, f"{L}{top}:{R}{top+1}", f"{p.name or '고객'}님", fill=NAVY, color=WHITE, size=13, bold=True)
    mid = (left + right) // 2
    spans = [(left, left+1), (left+2, mid), (mid+1, mid+2), (mid+3, right)]
    rows = [
        ("기존 월 보험료", won(p.old_monthly), "리모델링 후", won(p.after_monthly)),
        ("기존 납입 예정 총액", won(p.old_total), "변경 납입 예정 총액", won(p.after_total)),
    ]
    for rr, values in enumerate(rows, top+2):
        for idx, ((a, b), value) in enumerate(zip(spans, values)):
            _merge(ws, f"{get_column_letter(a)}{rr}:{get_column_letter(b)}{rr}", value,
                   fill=BLUE_LIGHT if idx % 2 == 0 else WHITE, color=NAVY2 if idx == 3 else INK,
                   size=7.5 if idx % 2 == 0 else 9.5, bold=True)
    row = top + 5
    _merge(ws, f"{L}{row}:{R}{row}", "새롭게 가입하는 보험", fill=NAVY2, color=WHITE, bold=True)
    text_end = right - 2
    _merge(ws, f"{L}{row+1}:{get_column_letter(text_end)}{row+1}", "보험 또는 보장 구성", fill=NAVY2, color=WHITE, size=8, bold=True)
    _merge(ws, f"{get_column_letter(text_end+1)}{row+1}:{R}{row+1}", "월 보험료", fill=NAVY2, color=WHITE, size=8, bold=True)
    shown = p.plans[:max_plans]
    for idx in range(max_plans):
        plan = shown[idx] if idx < len(shown) else NewPlan()
        rr = row + 2 + idx
        _merge(ws, f"{L}{rr}:{get_column_letter(text_end)}{rr}", plan.name, fill=SOFT, size=8, bold=bool(plan.name))
        _merge(ws, f"{get_column_letter(text_end+1)}{rr}:{R}{rr}", won(plan.monthly) if plan.name else "", fill=WHITE, size=8, bold=bool(plan.name))
    total_row = row + 2 + max_plans
    _merge(ws, f"{L}{total_row}:{get_column_letter(text_end)}{total_row}", "신규 보험료 합계", fill=GREEN_LIGHT, color=NAVY, bold=True)
    _merge(ws, f"{get_column_letter(text_end+1)}{total_row}:{R}{total_row}", won(p.new_plan_monthly), fill=GREEN_LIGHT, color=GREEN, bold=True)
    coverage_row = total_row + 2
    _merge(ws, f"{L}{coverage_row}:{R}{coverage_row}", "새롭게 확보되는 핵심 보장", fill=NAVY, color=WHITE, bold=True)
    _merge(ws, f"{L}{coverage_row+1}:{R}{coverage_row+2}", p.coverage, fill=GREEN_LIGHT, color=NAVY, size=8.5, bold=True)
    return coverage_row + 2


def _excel_new_panel(ws, p: Person, left: int, right: int, top: int, max_plans: int) -> int:
    L, R = get_column_letter(left), get_column_letter(right)
    _merge(ws, f"{L}{top}:{R}{top}", "새롭게 가입하는 보험", fill=NAVY, color=WHITE, bold=True)
    text_end = right - 2
    _merge(ws, f"{L}{top+1}:{get_column_letter(text_end)}{top+1}", "보험 또는 보장 구성", fill=NAVY2, color=WHITE, size=8, bold=True)
    _merge(ws, f"{get_column_letter(text_end+1)}{top+1}:{R}{top+1}", "월 보험료", fill=NAVY2, color=WHITE, size=8, bold=True)
    for idx in range(max_plans):
        plan = p.plans[idx] if idx < len(p.plans) else NewPlan()
        rr = top + 2 + idx
        _merge(ws, f"{L}{rr}:{get_column_letter(text_end)}{rr}", plan.name, fill=SOFT, size=8, bold=bool(plan.name))
        _merge(ws, f"{get_column_letter(text_end+1)}{rr}:{R}{rr}", won(plan.monthly) if plan.name else "", fill=WHITE, size=8, bold=bool(plan.name))
    total_row = top + 2 + max_plans
    _merge(ws, f"{L}{total_row}:{get_column_letter(text_end)}{total_row}", "신규 보험료 합계", fill=GREEN_LIGHT, color=NAVY, bold=True)
    _merge(ws, f"{get_column_letter(text_end+1)}{total_row}:{R}{total_row}", won(p.new_plan_monthly), fill=GREEN_LIGHT, color=GREEN, bold=True)
    coverage_row = total_row + 2
    _merge(ws, f"{L}{coverage_row}:{R}{coverage_row}", "새롭게 확보되는 핵심 보장", fill=NAVY, color=WHITE, bold=True)
    _merge(ws, f"{L}{coverage_row+1}:{R}{coverage_row+2}", p.coverage, fill=GREEN_LIGHT, color=NAVY, size=8.5, bold=True)
    return coverage_row + 2


def _excel_bottom(ws, people: list[Person], top: int) -> int:
    t = combined(people)
    spans = [("A", "C"), ("D", "H"), ("I", "M"), ("N", "R")]
    headers = ["한눈에 보는 비교" if len(people) == 1 else "2인 합산 비교", "기존", "리모델링 후", "변화"]
    for span, value in zip(spans, headers):
        _merge(ws, f"{span[0]}{top}:{span[1]}{top}", value, fill=NAVY, color=WHITE, bold=True)
    records = [
        ("월 보험료", t["old_monthly"], t["after_monthly"]),
        ("연간 보험료", t["old_monthly"]*12, t["after_monthly"]*12),
        ("납입 예정 총액", t["old_total"], t["after_total"]),
    ]
    for rr, (label, old, new) in enumerate(records, top+1):
        _merge(ws, f"A{rr}:C{rr}", label, fill=BLUE_LIGHT, bold=True)
        _merge(ws, f"D{rr}:H{rr}", won(old), fill=WHITE)
        _merge(ws, f"I{rr}:M{rr}", won(new), fill=WHITE)
        _merge(ws, f"N{rr}:P{rr}", change_amount(old, new), fill=GOLD_LIGHT, color=GREEN if old > new else INK, bold=True)
        _merge(ws, f"Q{rr}:R{rr}", change_rate(old, new), fill=GOLD_LIGHT, color=RED, bold=True)
    return top + 3


def _excel_detail(wb: Workbook, people: list[Person]) -> None:
    ws = wb.create_sheet("기존 계약 변경")
    ws.sheet_view.showGridLines = False
    for col, width in zip("ABCDEF", [16, 24, 18, 44, 16, 16]):
        ws.column_dimensions[col].width = width
    _merge(ws, "A1:F2", "기존 계약별 유지·감액·해지", color=NAVY, size=18, bold=True, border=False)
    row = 4
    for p in people:
        _merge(ws, f"A{row}:F{row}", f"{p.name or '고객'}님", fill=NAVY, color=WHITE, bold=True)
        row += 1
        for col, text in enumerate(["보험회사", "상품명", "처리 방향", "구체적인 변경 내용"], 1):
            end = col if col < 4 else 6
            _merge(ws, f"{get_column_letter(col)}{row}:{get_column_letter(end)}{row}", text, fill=NAVY2, color=WHITE, bold=True)
            if col == 4:
                break
        row += 1
        records = p.contracts or [ExistingContract(detail="입력된 기존 계약 변경 내용이 없습니다.")]
        for c in records:
            _merge(ws, f"A{row}:A{row}", c.company, fill=WHITE)
            _merge(ws, f"B{row}:B{row}", c.product, fill=WHITE)
            _merge(ws, f"C{row}:C{row}", c.action if c.company or c.product else "", fill=WHITE, bold=True)
            _merge(ws, f"D{row}:F{row}", c.detail, fill=WHITE)
            ws.row_dimensions[row].height = 30
            row += 1
        row += 1
    _merge(ws, f"A{row}:F{row}", "※ 감액·해지는 신규 계약의 승인과 보장 개시를 확인한 후 진행합니다.", fill=GOLD_LIGHT, color=MUTED, size=9)
    _excel_setup(ws, "F", row)


def create_excel(people: list[Person], title: str, detailed: bool, consultation_date: date, consultant: str) -> BytesIO:
    wb = Workbook()
    ws = wb.active
    ws.title = "리모델링 비교안"
    widths = [8,8,8,8,8,8,10,10,4,4,10,10,8,8,8,8,8,8]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    _excel_top(ws, people, title)
    if len(people) == 1:
        p = people[0]
        _excel_new_panel(ws, p, 11, 18, 8, 5)
        _merge(ws, "A8:H8", "보험료 비교", fill=NAVY, color=WHITE, bold=True)
        left_rows = [("기존 월 보험료", p.old_monthly), ("유지 보험료", p.retained_monthly), ("신규 보험료", p.new_plan_monthly), ("리모델링 후", p.after_monthly)]
        for rr, (label, value) in enumerate(left_rows, 9):
            _merge(ws, f"A{rr}:D{rr}", label, fill=BLUE_LIGHT, bold=True)
            _merge(ws, f"E{rr}:H{rr}", won(value), fill=GOLD_LIGHT if rr == 12 else WHITE, color=NAVY2, bold=True)
        _merge(ws, "A14:H14", "납입 예정 총액", fill=NAVY, color=WHITE, bold=True)
        for rr, (label, value) in enumerate([("기존", p.old_total), ("리모델링 후", p.after_total), ("변화", abs(p.total_change))], 15):
            _merge(ws, f"A{rr}:D{rr}", label, fill=BLUE_LIGHT, bold=True)
            _merge(ws, f"E{rr}:H{rr}", won(value), fill=GOLD_LIGHT if rr == 17 else WHITE, color=GREEN if rr == 17 and p.total_change < 0 else INK, bold=True)
        bottom = _excel_bottom(ws, people, 21)
        last = 27
    else:
        _excel_person_panel(ws, people[0], 1, 8, 8, 4)
        _excel_person_panel(ws, people[1], 11, 18, 8, 4)
        bottom = _excel_bottom(ws, people, 25)
        last = 31
    excluded = [p.name for person in people for p in person.excluded_plans]
    note = "※ 갱신형·종신납·계속 납입 보험료는 월 보험료에는 포함되며 납입 예정 총액 비교에서는 제외됩니다."
    _merge(ws, f"A{last-1}:R{last-1}", note, color=MUTED, size=8, border=False)
    _merge(ws, f"A{last}:R{last}", f"상담일 {consultation_date:%Y.%m.%d} · 담당자 {consultant or '-'}", color=MUTED, size=8, border=False)
    _excel_setup(ws, "R", last)
    if detailed:
        _excel_detail(wb, people)
    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out


# ---------------- Word ----------------
def _wfont(run, size: float, color: str = INK, bold: bool = False) -> None:
    run.font.name = "맑은 고딕"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _wshade(cell, color: str) -> None:
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd")) or OxmlElement("w:shd")
    if shd.getparent() is None:
        pr.append(shd)
    shd.set(qn("w:fill"), color)


def _wcell(cell, text: str = "", *, fill: str = WHITE, color: str = INK, size: float = 9, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _wshade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _wfont(run, size, color, bold)


def _wtable(doc_or_cell, rows: int, cols: int, widths: list[float] | None = None):
    table = doc_or_cell.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def _wchange_cell(cell, old: int, new: int, size: float = 9) -> None:
    _wcell(cell, "", fill=GOLD_LIGHT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(change_amount(old, new))
    _wfont(r1, size, GREEN if old > new else INK, True)
    pct = change_rate(old, new)
    if pct:
        r2 = p.add_run(f" · {pct}")
        _wfont(r2, size, RED, True)


def _word_top(doc: Document, people: list[Person], title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _wfont(p.add_run(title), 19, NAVY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    _wfont(p.add_run("보험료 부담과 필요한 보장을 함께 비교해 오래 유지하기 쉬운 구조로 재구성했습니다."), 9, MUTED, True)
    t = combined(people)
    labels = (["월 보험료 변화", "변경 후 월 보험료", "납입 예정 총액 변화"] if len(people) == 1 else
              ["합산 월 보험료 변화", "변경 후 합산 월 보험료", "납입 예정 총액 변화"])
    table = _wtable(doc, 2, 3, [9.3, 9.3, 9.3])
    for i, label in enumerate(labels):
        _wcell(table.cell(0, i), label, fill=NAVY2, color=WHITE, size=9, bold=True)
    _wchange_cell(table.cell(1, 0), t["old_monthly"], t["after_monthly"], 9.5)
    _wcell(table.cell(1, 1), won(t["after_monthly"]), fill=GOLD_LIGHT, color=NAVY, size=10, bold=True)
    _wchange_cell(table.cell(1, 2), t["old_total"], t["after_total"], 9.5)


def _word_person(cell, p: Person, max_plans: int) -> None:
    cell.text = ""
    _wcell(_wtable(cell, 1, 1).cell(0, 0), f"{p.name or '고객'}님", fill=NAVY, color=WHITE, size=12, bold=True)
    premium = _wtable(cell, 2, 4, [2.3, 3.5, 2.3, 3.5])
    rows = [("기존 월 보험료", won(p.old_monthly), "리모델링 후", won(p.after_monthly)),
            ("기존 납입 예정 총액", won(p.old_total), "변경 납입 예정 총액", won(p.after_total))]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            _wcell(premium.cell(r, c), value, fill=BLUE_LIGHT if c % 2 == 0 else WHITE, color=NAVY2 if c == 3 else INK, size=7.5 if c % 2 == 0 else 9, bold=True)
    plans = _wtable(cell, max_plans+2, 2, [8.2, 3.4])
    _wcell(plans.cell(0, 0), "보험 또는 보장 구성", fill=NAVY2, color=WHITE, size=8, bold=True)
    _wcell(plans.cell(0, 1), "월 보험료", fill=NAVY2, color=WHITE, size=8, bold=True)
    for r in range(max_plans):
        plan = p.plans[r] if r < len(p.plans) else NewPlan()
        _wcell(plans.cell(r+1, 0), plan.name, fill=SOFT, size=8, bold=bool(plan.name))
        _wcell(plans.cell(r+1, 1), won(plan.monthly) if plan.name else "", size=8, bold=bool(plan.name))
    _wcell(plans.cell(max_plans+1, 0), "신규 보험료 합계", fill=GREEN_LIGHT, color=NAVY, size=8.5, bold=True)
    _wcell(plans.cell(max_plans+1, 1), won(p.new_plan_monthly), fill=GREEN_LIGHT, color=GREEN, size=8.5, bold=True)
    coverage = _wtable(cell, 2, 1)
    _wcell(coverage.cell(0, 0), "새롭게 확보되는 핵심 보장", fill=NAVY, color=WHITE, size=8.5, bold=True)
    _wcell(coverage.cell(1, 0), p.coverage, fill=GREEN_LIGHT, color=NAVY, size=8.5, bold=True)


def _word_new_panel(cell, p: Person, max_plans: int) -> None:
    cell.text = ""
    title = _wtable(cell, 1, 1)
    _wcell(title.cell(0, 0), "새롭게 가입하는 보험", fill=NAVY, color=WHITE, size=9, bold=True)
    plans = _wtable(cell, max_plans + 2, 2, [11.5, 4.9])
    _wcell(plans.cell(0, 0), "보험 또는 보장 구성", fill=NAVY2, color=WHITE, size=8, bold=True)
    _wcell(plans.cell(0, 1), "월 보험료", fill=NAVY2, color=WHITE, size=8, bold=True)
    for r in range(max_plans):
        plan = p.plans[r] if r < len(p.plans) else NewPlan()
        _wcell(plans.cell(r+1, 0), plan.name, fill=SOFT, size=8, bold=bool(plan.name))
        _wcell(plans.cell(r+1, 1), won(plan.monthly) if plan.name else "", size=8, bold=bool(plan.name))
    _wcell(plans.cell(max_plans+1, 0), "신규 보험료 합계", fill=GREEN_LIGHT, color=NAVY, size=8.5, bold=True)
    _wcell(plans.cell(max_plans+1, 1), won(p.new_plan_monthly), fill=GREEN_LIGHT, color=GREEN, size=8.5, bold=True)
    coverage = _wtable(cell, 2, 1)
    _wcell(coverage.cell(0, 0), "새롭게 확보되는 핵심 보장", fill=NAVY, color=WHITE, size=8.5, bold=True)
    _wcell(coverage.cell(1, 0), p.coverage, fill=GREEN_LIGHT, color=NAVY, size=8.5, bold=True)


def _word_bottom(doc: Document, people: list[Person]) -> None:
    t = combined(people)
    table = _wtable(doc, 4, 4, [4.2, 7.4, 7.4, 8.9])
    headers = ["한눈에 보는 비교" if len(people) == 1 else "2인 합산 비교", "기존", "리모델링 후", "변화"]
    for c, value in enumerate(headers):
        _wcell(table.cell(0, c), value, fill=NAVY, color=WHITE, size=8.5, bold=True)
    records = [("월 보험료", t["old_monthly"], t["after_monthly"]),
               ("연간 보험료", t["old_monthly"]*12, t["after_monthly"]*12),
               ("납입 예정 총액", t["old_total"], t["after_total"])]
    for r, (label, old, new) in enumerate(records, 1):
        _wcell(table.cell(r, 0), label, fill=BLUE_LIGHT, size=8.5, bold=True)
        _wcell(table.cell(r, 1), won(old), size=8.5)
        _wcell(table.cell(r, 2), won(new), size=8.5)
        _wchange_cell(table.cell(r, 3), old, new, 8.3)


def create_word(people: list[Person], title: str, detailed: bool, consultation_date: date, consultant: str) -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Cm(29.7), Cm(21.0)
    section.left_margin = section.right_margin = Cm(.75)
    section.top_margin = section.bottom_margin = Cm(.55)
    _word_top(doc, people, title)
    outer = _wtable(doc, 1, 3, [13.5, .7, 13.5] if len(people) == 2 else [10.5, .8, 16.4])
    if len(people) == 2:
        _word_person(outer.cell(0, 0), people[0], 4)
        _word_person(outer.cell(0, 2), people[1], 4)
    else:
        p = people[0]
        left = outer.cell(0, 0)
        left.text = ""
        compare = _wtable(left, 9, 2, [5.1, 5.1])
        _wcell(compare.cell(0, 0), "보험료 비교", fill=NAVY, color=WHITE, bold=True)
        compare.cell(0, 0).merge(compare.cell(0, 1))
        rows = [("기존 월 보험료", p.old_monthly), ("유지 보험료", p.retained_monthly), ("신규 보험료", p.new_plan_monthly), ("리모델링 후", p.after_monthly),
                ("납입 예정 총액", 0), ("기존", p.old_total), ("리모델링 후", p.after_total), ("변화", abs(p.total_change))]
        for r, (label, value) in enumerate(rows, 1):
            _wcell(compare.cell(r, 0), label, fill=BLUE_LIGHT, size=8, bold=True)
            _wcell(compare.cell(r, 1), won(value), fill=GOLD_LIGHT if label == "변화" else WHITE, color=GREEN if label == "변화" and p.total_change < 0 else INK, size=8.5, bold=True)
        _word_new_panel(outer.cell(0, 2), p, 5)
    _word_bottom(doc, people)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    _wfont(p.add_run("※ 갱신형·종신납·계속 납입 보험료는 월 보험료에는 포함되며 납입 예정 총액 비교에서는 제외됩니다."), 7.5, MUTED)
    _wfont(p.add_run(f"    상담일 {consultation_date:%Y.%m.%d} · 담당자 {consultant or '-'}"), 7.5, MUTED)
    if detailed:
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _wfont(p.add_run("기존 계약별 유지·감액·해지"), 18, NAVY, True)
        for person in people:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _wfont(p.add_run(f"{person.name or '고객'}님"), 11, NAVY, True)
            table = _wtable(doc, max(2, len(person.contracts)+1), 4, [4, 6, 5, 13])
            for c, text in enumerate(["보험회사", "상품명", "처리 방향", "구체적인 변경 내용"]):
                _wcell(table.cell(0, c), text, fill=NAVY2, color=WHITE, size=8.5, bold=True)
            records = person.contracts or [ExistingContract(detail="입력된 기존 계약 변경 내용이 없습니다.")]
            for r, item in enumerate(records, 1):
                for c, value in enumerate([item.company, item.product, item.action if item.company or item.product else "", item.detail]):
                    _wcell(table.cell(r, c), value, size=8.5, bold=c == 2)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def load_example(count: int) -> None:
    examples = [
        ("유민재", "694580", "144940000", "0", "0", [("암·뇌·심장 진단비", "128589"), ("암 주요치료비", "111255"), ("운전자보험", "15000"), ("순환계 주요치료비 및 추가 보장", "112863")]),
        ("김래아", "495470", "102950000", "0", "0", [("암·뇌·심장 진단비", "95691"), ("암 주요치료비", "101245"), ("운전자보험", "15000"), ("순환계 주요치료비 및 추가 보장", "117274")]),
    ]
    for idx in range(count):
        no = idx + 1
        name, oldm, oldt, keepm, keept, plans = examples[idx]
        st.session_state[f"rm_name_{no}"] = name
        st.session_state[f"rm_old_monthly_{no}"] = oldm
        st.session_state[f"rm_old_total_{no}"] = oldt
        st.session_state[f"rm_retained_monthly_{no}"] = keepm
        st.session_state[f"rm_retained_total_{no}"] = keept
        st.session_state[f"rm_plan_count_{no}"] = 4
        for i, (pname, premium) in enumerate(plans):
            st.session_state[f"rm_plan_name_{no}_{i}"] = pname
            st.session_state[f"rm_plan_premium_{no}_{i}"] = premium
            st.session_state[f"rm_plan_type_{no}_{i}"] = "확정 납입" if "운전자" not in pname else "계속 납입"
            st.session_state[f"rm_plan_years_{no}_{i}"] = 20
        st.session_state[f"rm_coverage_{no}"] = "암·뇌·심장 진단비 보완 · 주요 치료비 강화 · 보장 공백 보완"


def run() -> None:
    page_header("고객 상담", APP_TITLE, "간편 입력으로 한눈에 보는 비교표를 만들고 워드·엑셀로 내려받습니다.", "RM")
    c1, c2, c3 = st.columns(3)
    with c1:
        count = int(st.selectbox("대상 인원", [1, 2], format_func=lambda x: f"{x}명", key="rm_count"))
    with c2:
        mode = st.radio("입력 모드", ["간편 모드", "상세 모드"], horizontal=True, key="rm_mode")
    with c3:
        if st.button("예시 데이터 입력", use_container_width=True):
            load_example(count)
            st.rerun()
    detailed = mode == "상세 모드"
    meta1, meta2 = st.columns(2)
    with meta1:
        consultation_date = st.date_input("상담일", value=date.today(), key="rm_date")
    with meta2:
        consultant = st.text_input("담당자", key="rm_consultant", placeholder="예: 박병선 팀장")
    people: list[Person] = []
    tabs = st.tabs([f"고객 {i}" for i in range(1, count+1)])
    for i, tab in enumerate(tabs, 1):
        with tab:
            people.append(render_person_inputs(i, detailed))
    if count == 2:
        shared = st.checkbox("두 고객의 핵심 보장을 하나로 묶어 표시", key="rm_shared_coverage")
        if shared:
            shared_text = st.text_area("공통 핵심 보장", key="rm_shared_coverage_text", height=75)
            for p in people:
                p.coverage = clean(shared_text)
    names = [p.name for p in people if p.name]
    default_title = " · ".join(f"{n}님" for n in names) + " 보험 리모델링 비교안" if names else "보험 리모델링 비교안"
    title = st.text_input("자료 제목", value=default_title, key="rm_title")
    st.divider()
    st.subheader("자동 계산 결과")
    t = combined(people)
    a, b, c = st.columns(3)
    labels = (["월 보험료 변화", "변경 후 월 보험료", "납입 예정 총액 변화"] if count == 1 else
              ["합산 월 보험료 변화", "변경 후 합산 월 보험료", "납입 예정 총액 변화"])
    a.metric(labels[0], change_amount(t["old_monthly"], t["after_monthly"]), change_rate(t["old_monthly"], t["after_monthly"]) or None)
    b.metric(labels[1], won(t["after_monthly"]))
    c.metric(labels[2], change_amount(t["old_total"], t["after_total"]), change_rate(t["old_total"], t["after_total"]) or None)
    missing = [f"고객 {i+1} 이름" for i, p in enumerate(people) if not p.name]
    if missing:
        st.warning("입력 필요: " + ", ".join(missing))
        return
    excel = create_excel(people, clean(title) or default_title, detailed, consultation_date, clean(consultant))
    word = create_word(people, clean(title) or default_title, detailed, consultation_date, clean(consultant))
    base = safe_filename(f"{clean(title) or default_title}_{consultation_date:%Y%m%d}")
    st.markdown("#### 다운로드 형식 선택")
    d1, d2 = st.columns(2)
    with d1:
        st.download_button("엑셀로 다운로드", excel, f"{base}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    with d2:
        st.download_button("워드로 다운로드", word, f"{base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")


if __name__ == "__main__":
    run()
